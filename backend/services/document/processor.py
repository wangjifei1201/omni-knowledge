"""
Document Processing Service
Handles document parsing, chunking, and indexing
"""

import asyncio
import json
import os
import uuid
from typing import Optional

from core.config import get_settings
from loguru import logger
from services.document.chunking_strategy import ChunkResult, chunking_engine
from services.rag.vector_store import faiss_vector_store
from services.storage.local_storage import local_storage

settings = get_settings()


class DocumentProcessor:
    """
    Document processing pipeline:
    1. Read document content from storage
    2. Parse document based on file type
    3. Split content into chunks (using configurable strategy)
    4. Persist chunks to database
    5. Generate embeddings and store in vector database
    """

    def __init__(self):
        pass

    async def process_document(
        self,
        doc_id: str,
        file_path: str,
        file_type: str,
        doc_name: str,
        chunking_strategy: str = "paragraph",
        chunking_params: Optional[dict] = None,
    ) -> dict:
        """
        Process a document: parse, chunk, persist, and index.

        Returns:
            dict with processing results: chunk_count, page_count, status
        """
        try:
            logger.info(f"Processing document: {doc_id} ({doc_name}), strategy={chunking_strategy}")

            # 1. Read document content
            content = await local_storage.download_file(file_path)
            if not content:
                return {"status": "failed", "error": "文件不存在", "chunk_count": 0}

            # 2. Parse document based on type
            text_content, page_count = await self._parse_document(content, file_type, doc_name)

            if not text_content:
                return {"status": "failed", "error": "无法解析文档内容", "chunk_count": 0}

            # 3. Split into chunks using the chunking engine
            chunk_results = chunking_engine.chunk_document(
                text=text_content,
                doc_id=doc_id,
                doc_name=doc_name,
                strategy=chunking_strategy,
                params=chunking_params,
            )

            if not chunk_results:
                return {"status": "completed", "chunk_count": 0, "page_count": page_count}

            # 4. Persist chunks to database
            await self._persist_chunks_to_db(chunk_results, doc_id)

            # 5. Generate embeddings and store in vector database
            await self._index_chunks(chunk_results, doc_id, doc_name)

            logger.info(f"Document processed: {doc_id}, {len(chunk_results)} chunks indexed")

            return {
                "status": "completed",
                "chunk_count": len(chunk_results),
                "page_count": page_count,
            }

        except Exception as e:
            logger.error(f"Error processing document {doc_id}: {e}")
            return {"status": "failed", "error": str(e), "chunk_count": 0}

    async def _persist_chunks_to_db(self, chunks: list[ChunkResult], doc_id: str):
        """Persist chunk results to the DocumentChunk database table"""
        from core.database import AsyncSessionLocal
        from models.document import DocumentChunk
        from sqlalchemy import delete

        async with AsyncSessionLocal() as db:
            # Delete existing chunks for this document
            await db.execute(delete(DocumentChunk).where(DocumentChunk.doc_id == doc_id))

            # Batch insert new chunks
            chunk_records = []
            for chunk in chunks:
                record = DocumentChunk(
                    doc_id=doc_id,
                    chunk_index=chunk.chunk_index,
                    chunk_type="text",
                    content=chunk.content,
                    token_count=chunk.token_count,
                    metadata_extra=json.dumps(chunk.metadata, ensure_ascii=False) if chunk.metadata else None,
                )
                chunk_records.append(record)

            db.add_all(chunk_records)
            await db.commit()
            logger.info(f"Persisted {len(chunk_records)} chunks to database for doc {doc_id}")

    async def _parse_document(
        self,
        content: bytes,
        file_type: str,
        doc_name: str,
    ) -> tuple[str, int]:
        """
        Parse document content based on file type.
        Returns (text_content, page_count)
        """
        try:
            if file_type in ["txt", "md", "csv"]:
                # Plain text files
                text = content.decode("utf-8", errors="ignore")
                return text, 1

            elif file_type == "pdf":
                # Try to parse PDF
                return await self._parse_pdf(content)

            elif file_type in ["doc", "docx"]:
                # Try to parse Word
                return await self._parse_word(content, file_type)

            elif file_type in ["xls", "xlsx"]:
                # Try to parse Excel
                return await self._parse_excel(content, file_type)

            elif file_type in ["png", "jpg", "jpeg", "gif", "bmp"]:
                # Images - placeholder for OCR
                return f"[图片文件: {doc_name}]", 1

            else:
                # Unknown type - try as text
                try:
                    text = content.decode("utf-8", errors="ignore")
                    return text, 1
                except:
                    return "", 0

        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            return "", 0

    async def _parse_pdf(self, content: bytes) -> tuple[str, int]:
        """Parse PDF file"""
        try:
            import io

            import pdfplumber

            text_parts = []
            page_count = 0

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return "\n\n".join(text_parts), page_count

        except ImportError:
            logger.warning("pdfplumber not installed, cannot parse PDF")
            return "[PDF文件，需要安装pdfplumber解析]", 1
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return "", 0

    async def _parse_word(self, content: bytes, file_type: str) -> tuple[str, int]:
        """Parse Word document (.doc and .docx)"""
        try:
            if file_type == "docx":
                import io

                from docx import Document

                doc = Document(io.BytesIO(content))
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)

                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)

                return "\n\n".join(text_parts), 1
            else:
                # .doc format (Word 97-2003)
                return await self._parse_doc_legacy(content)

        except ImportError as e:
            logger.warning(f"Required library not installed: {e}")
            return "[Word文件，需要安装python-docx解析]", 1
        except Exception as e:
            logger.error(f"Error parsing Word: {e}")
            return "", 0

    async def _parse_doc_legacy(self, content: bytes) -> tuple[str, int]:
        """Parse legacy .doc format (Word 97-2003) using olefile"""
        import io
        import re
        import struct

        # Try olefile to extract text from Word binary format
        try:
            import olefile

            ole = olefile.OleFileIO(io.BytesIO(content))

            # Read WordDocument stream which contains the text
            if ole.exists("WordDocument"):
                word_doc = ole.openstream("WordDocument").read()

                # Read 1Table or 0Table stream for text positions
                table_stream = None
                if ole.exists("1Table"):
                    table_stream = ole.openstream("1Table").read()
                elif ole.exists("0Table"):
                    table_stream = ole.openstream("0Table").read()

                # Extract text using simple method - find text in WordDocument
                # Word 97-2003 stores text as bytes that can be decoded
                text = self._extract_text_from_word_doc(word_doc)

                ole.close()

                if text and len(text) > 50:
                    logger.info(f"Parsed .doc using olefile, extracted {len(text)} chars")
                    return text, 1

            ole.close()

        except ImportError:
            logger.debug("olefile not installed")
        except Exception as e:
            logger.debug(f"olefile extraction failed: {e}")

        # Fallback: try to extract any readable text from the binary
        try:
            text = self._extract_text_from_binary(content)
            if text and len(text) > 100:
                logger.info(f"Parsed .doc using binary extraction, extracted {len(text)} chars")
                return text, 1
        except Exception as e:
            logger.debug(f"Binary extraction failed: {e}")

        logger.warning("Could not parse .doc file - please convert to .docx format")
        return "[DOC文件解析失败，建议将文件另存为.docx格式后重新上传]", 1

    def _extract_text_from_word_doc(self, word_doc: bytes) -> str:
        """Extract text from WordDocument stream"""
        import re

        # Try multiple encodings
        text_parts = []

        # Method 1: Look for text between common markers
        # Word documents often have text in specific byte patterns

        # Try to find continuous text blocks
        # Word 97-2003 often stores text as CP1252 or UTF-16
        for encoding in ["cp1252", "utf-16-le", "gb2312", "gbk", "utf-8"]:
            try:
                decoded = word_doc.decode(encoding, errors="ignore")
                # Filter to keep only printable chars
                filtered = "".join(c for c in decoded if c.isprintable() or c in "\n\r\t ")
                # Remove excessive whitespace
                filtered = re.sub(r"\s+", " ", filtered)
                # Find chunks of meaningful text (Chinese or ASCII)
                chunks = re.findall(
                    r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffefa-zA-Z0-9，。！？、；：" "''（）【】《》\\s]{20,}", filtered
                )
                if chunks:
                    text_parts.extend(chunks)
            except:
                continue

        if text_parts:
            # Deduplicate and join
            seen = set()
            unique_parts = []
            for part in text_parts:
                part = part.strip()
                if part and part not in seen and len(part) > 20:
                    seen.add(part)
                    unique_parts.append(part)

            return "\n\n".join(unique_parts)

        return ""

    def _extract_text_from_binary(self, content: bytes) -> str:
        """Extract readable text from binary content as fallback"""
        import re

        text_parts = []

        # Try different encodings
        for encoding in ["gb2312", "gbk", "cp1252", "utf-8"]:
            try:
                decoded = content.decode(encoding, errors="ignore")
                # Find Chinese text blocks
                chunks = re.findall(r"[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef，。！？、；：" "''（）【】《》]{10,}", decoded)
                text_parts.extend(chunks)
            except:
                continue

        if text_parts:
            # Sort by length and take longer chunks
            text_parts = sorted(set(text_parts), key=len, reverse=True)
            # Filter out very short or duplicate chunks
            result = []
            seen_content = set()
            for part in text_parts:
                # Check if this is a substring of existing content
                is_duplicate = False
                for seen in seen_content:
                    if part in seen:
                        is_duplicate = True
                        break
                if not is_duplicate:
                    result.append(part)
                    seen_content.add(part)

            return "\n\n".join(result[:50])  # Limit to 50 chunks

        return ""

    async def _parse_excel(self, content: bytes, file_type: str) -> tuple[str, int]:
        """Parse Excel file"""
        try:
            import io

            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(io.BytesIO(content))
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"=== 工作表: {sheet_name} ===")
                # Convert to markdown-like table
                text_parts.append(df.to_string())

            return "\n\n".join(text_parts), len(excel_file.sheet_names)

        except ImportError:
            logger.warning("pandas/openpyxl not installed")
            return "[Excel文件，需要安装pandas和openpyxl解析]", 1
        except Exception as e:
            logger.error(f"Error parsing Excel: {e}")
            return "", 0

    async def _index_chunks(
        self,
        chunks: list[ChunkResult],
        doc_id: str,
        doc_name: str,
    ):
        """
        Generate embeddings and store chunks in vector database.
        Uses real embedding service to generate vectors.
        """
        from services.llm.embedding import embedding_service

        if not chunks:
            return

        # Extract chunk contents for batch embedding
        chunk_contents = [chunk.content for chunk in chunks]

        # Generate real embeddings using embedding service
        logger.info(f"Generating embeddings for {len(chunks)} chunks...")
        vectors = await embedding_service.embed_batch(chunk_contents)

        # Filter out chunks with failed embeddings
        valid_vectors = []
        chunk_ids = []
        metadatas = []

        for i, (chunk, vector) in enumerate(zip(chunks, vectors)):
            if vector and len(vector) == settings.EMBEDDING_DIMENSION:
                chunk_id = f"{doc_id}_chunk_{chunk.chunk_index}"
                valid_vectors.append(vector)
                chunk_ids.append(chunk_id)
                metadatas.append(
                    {
                        "doc_id": doc_id,
                        "doc_name": doc_name,
                        "content": chunk.content,
                        "chunk_index": chunk.chunk_index,
                    }
                )
            else:
                logger.warning(f"Skipping chunk {i} - embedding failed")

        if valid_vectors:
            # Store in FAISS
            await faiss_vector_store.add_vectors(valid_vectors, chunk_ids, metadatas)
            logger.info(f"Indexed {len(valid_vectors)} chunks into vector store")


# Singleton instance
document_processor = DocumentProcessor()
